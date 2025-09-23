"""
Content Processing Pipeline
AI Language Tutor App - YouLearn Functionality Implementation

Provides:
- YouTube video processing and transcript extraction
- Document parsing (PDF, DOCX, text files)
- Learning material generation (flashcards, quizzes, summaries)
- Content organization and library management
- Real-time processing status updates

Target: Process YouTube videos → learning materials in <2 minutes
"""

import asyncio
import logging
import tempfile
import hashlib
import json
from datetime import datetime
from typing import Dict, List, Any, Optional, Union, AsyncGenerator
from pathlib import Path
from dataclasses import dataclass, asdict
from enum import Enum
import aiohttp
import aiofiles
from urllib.parse import urlparse, parse_qs

# Document processing
import PyPDF2
from docx import Document

# YouTube processing
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
import yt_dlp

# AI Services
from app.services.ai_router import ai_router, generate_ai_response
from app.core.config import get_settings

logger = logging.getLogger(__name__)


class ContentType(Enum):
    """Content type classifications"""

    YOUTUBE_VIDEO = "youtube_video"
    PDF_DOCUMENT = "pdf_document"
    WORD_DOCUMENT = "word_document"
    TEXT_FILE = "text_file"
    WEB_ARTICLE = "web_article"
    AUDIO_FILE = "audio_file"
    IMAGE_FILE = "image_file"
    UNKNOWN = "unknown"


class ProcessingStatus(Enum):
    """Content processing status"""

    QUEUED = "queued"
    EXTRACTING = "extracting"
    ANALYZING = "analyzing"
    GENERATING = "generating"
    ORGANIZING = "organizing"
    COMPLETED = "completed"
    FAILED = "failed"


class LearningMaterialType(Enum):
    """Types of learning materials that can be generated"""

    SUMMARY = "summary"
    FLASHCARDS = "flashcards"
    QUIZ = "quiz"
    NOTES = "notes"
    MIND_MAP = "mind_map"
    KEY_CONCEPTS = "key_concepts"
    PRACTICE_QUESTIONS = "practice_questions"


@dataclass
class ProcessingProgress:
    """Real-time processing progress tracking"""

    content_id: str
    status: ProcessingStatus
    current_step: str
    progress_percentage: int
    time_elapsed: float
    estimated_remaining: float
    details: str
    error_message: Optional[str] = None
    created_at: datetime = None

    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now()


@dataclass
class ContentMetadata:
    """Metadata for processed content"""

    content_id: str
    title: str
    content_type: ContentType
    source_url: Optional[str]
    language: str
    duration: Optional[float]  # In minutes
    word_count: int
    difficulty_level: str
    topics: List[str]
    author: Optional[str]
    created_at: datetime
    file_size: Optional[int] = None

    def __post_init__(self):
        if isinstance(self.content_type, str):
            self.content_type = ContentType(self.content_type)


@dataclass
class LearningMaterial:
    """Generated learning material"""

    material_id: str
    content_id: str
    material_type: LearningMaterialType
    title: str
    content: Dict[str, Any]  # Structure depends on material type
    difficulty_level: str
    estimated_time: int  # In minutes
    tags: List[str]
    created_at: datetime = None

    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now()
        if isinstance(self.material_type, str):
            self.material_type = LearningMaterialType(self.material_type)


@dataclass
class ProcessedContent:
    """Complete processed content with metadata and materials"""

    metadata: ContentMetadata
    raw_content: str
    processed_content: str
    learning_materials: List[LearningMaterial]
    processing_stats: Dict[str, Any]


class ContentProcessor:
    """Main content processing service implementing YouLearn functionality"""

    def __init__(self):
        self.settings = get_settings()
        self.processing_progress: Dict[str, ProcessingProgress] = {}
        self.content_library: Dict[str, ProcessedContent] = {}
        self.temp_dir = Path(tempfile.gettempdir()) / "ai_tutor_content"
        self.temp_dir.mkdir(exist_ok=True)

        # Processing configuration
        self.max_content_length = 50000  # Maximum characters to process
        self.supported_languages = [
            "en",
            "es",
            "fr",
            "zh",
            "de",
            "it",
            "pt",
            "ja",
            "ko",
        ]
        self.processing_timeout = 120  # 2 minutes maximum

    def _generate_content_id(self, source: str) -> str:
        """Generate unique content ID from source"""
        timestamp = datetime.now().isoformat()
        unique_string = f"{source}_{timestamp}"
        return hashlib.md5(unique_string.encode()).hexdigest()[:12]

    def _update_progress(
        self,
        content_id: str,
        status: ProcessingStatus,
        step: str,
        percentage: int,
        details: str = "",
        error: Optional[str] = None,
    ):
        """Update processing progress"""
        if content_id in self.processing_progress:
            progress = self.processing_progress[content_id]
            progress.status = status
            progress.current_step = step
            progress.progress_percentage = percentage
            progress.time_elapsed = (
                datetime.now() - progress.created_at
            ).total_seconds()
            progress.details = details
            progress.error_message = error

            # Estimate remaining time
            if percentage > 0 and percentage < 100:
                progress.estimated_remaining = (progress.time_elapsed / percentage) * (
                    100 - percentage
                )
            else:
                progress.estimated_remaining = 0
        else:
            self.processing_progress[content_id] = ProcessingProgress(
                content_id=content_id,
                status=status,
                current_step=step,
                progress_percentage=percentage,
                time_elapsed=0,
                estimated_remaining=0,
                details=details,
                error_message=error,
            )

    async def get_processing_progress(
        self, content_id: str
    ) -> Optional[ProcessingProgress]:
        """Get current processing progress for content"""
        return self.processing_progress.get(content_id)

    def _detect_content_type(
        self, source: str, file_path: Optional[Path] = None
    ) -> ContentType:
        """Detect content type from source URL or file"""

        # YouTube detection
        youtube_domains = ["youtube.com", "youtu.be", "m.youtube.com"]
        if any(domain in source.lower() for domain in youtube_domains):
            return ContentType.YOUTUBE_VIDEO

        # File extension detection
        if file_path:
            extension = file_path.suffix.lower()
            if extension == ".pdf":
                return ContentType.PDF_DOCUMENT
            elif extension in [".docx", ".doc"]:
                return ContentType.WORD_DOCUMENT
            elif extension in [".txt", ".md", ".rtf"]:
                return ContentType.TEXT_FILE
            elif extension in [".mp3", ".wav", ".m4a", ".flac"]:
                return ContentType.AUDIO_FILE
            elif extension in [".jpg", ".jpeg", ".png", ".gif", ".bmp"]:
                return ContentType.IMAGE_FILE

        # URL detection
        if source.startswith(("http://", "https://")):
            return ContentType.WEB_ARTICLE

        return ContentType.UNKNOWN

    def _extract_youtube_id(self, url: str) -> Optional[str]:
        """Extract YouTube video ID from URL"""
        try:
            # Parse URL
            parsed = urlparse(url)

            # Different YouTube URL formats
            if parsed.hostname in ["youtu.be"]:
                return parsed.path[1:]  # Remove leading slash
            elif parsed.hostname in ["youtube.com", "www.youtube.com", "m.youtube.com"]:
                if "watch" in parsed.path:
                    query_params = parse_qs(parsed.query)
                    return query_params.get("v", [None])[0]
                elif "embed" in parsed.path:
                    return parsed.path.split("/")[-1]

            return None
        except Exception as e:
            logger.error(f"Failed to extract YouTube ID from {url}: {e}")
            return None

    async def _extract_youtube_content(self, url: str) -> Dict[str, Any]:
        """Extract content from YouTube video"""
        video_id = self._extract_youtube_id(url)
        if not video_id:
            raise ValueError(f"Invalid YouTube URL: {url}")

        try:
            # Get video metadata using yt-dlp
            ydl_opts = {"quiet": True, "no_warnings": True, "extract_flat": False}

            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)

                title = info.get("title", "Unknown Title")
                description = info.get("description", "")
                duration = (
                    info.get("duration", 0) / 60 if info.get("duration") else 0
                )  # Convert to minutes
                uploader = info.get("uploader", "Unknown")

            # Get transcript
            try:
                # Try to get transcript in multiple languages
                transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

                # Prefer English, then any available language
                transcript = None
                for lang_code in ["en", "en-US", "en-GB"]:
                    try:
                        transcript = transcript_list.find_transcript([lang_code])
                        break
                    except:
                        continue

                if not transcript:
                    # Get any available transcript
                    transcript = (
                        transcript_list.find_generated_transcripts()[0]
                        if transcript_list.find_generated_transcripts()
                        else None
                    )

                if transcript:
                    transcript_data = transcript.fetch()
                    formatter = TextFormatter()
                    transcript_text = formatter.format_transcript(transcript_data)
                else:
                    transcript_text = description  # Fallback to description

            except Exception as e:
                logger.warning(f"Could not get transcript for {video_id}: {e}")
                transcript_text = description  # Fallback to description

            return {
                "title": title,
                "content": transcript_text,
                "description": description,
                "duration": duration,
                "author": uploader,
                "language": "en",  # Default to English
                "word_count": len(transcript_text.split()),
            }

        except Exception as e:
            logger.error(f"Failed to extract YouTube content from {url}: {e}")
            raise ValueError(f"Could not process YouTube video: {str(e)}")

    async def _extract_pdf_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from PDF file"""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Extract text from all pages
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"

                # Get metadata
                info = pdf_reader.metadata
                title = info.get("/Title", file_path.stem) if info else file_path.stem
                author = info.get("/Author", "Unknown") if info else "Unknown"

                return {
                    "title": title,
                    "content": text_content.strip(),
                    "author": author,
                    "language": "en",  # TODO: Detect language
                    "word_count": len(text_content.split()),
                    "page_count": len(pdf_reader.pages),
                }

        except Exception as e:
            logger.error(f"Failed to extract PDF content from {file_path}: {e}")
            raise ValueError(f"Could not process PDF file: {str(e)}")

    async def _extract_docx_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from Word document"""
        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            # Try to get title from first paragraph or filename
            title = doc.paragraphs[0].text[:100] if doc.paragraphs else file_path.stem
            if len(title) > 50:
                title = title[:50] + "..."

            return {
                "title": title,
                "content": text_content.strip(),
                "author": "Unknown",
                "language": "en",  # TODO: Detect language
                "word_count": len(text_content.split()),
                "paragraph_count": len(doc.paragraphs),
            }

        except Exception as e:
            logger.error(f"Failed to extract DOCX content from {file_path}: {e}")
            raise ValueError(f"Could not process Word document: {str(e)}")

    async def _extract_text_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from text file"""
        try:
            async with aiofiles.open(file_path, "r", encoding="utf-8") as file:
                content = await file.read()

            # Use first line as title, or filename
            lines = content.split("\n")
            title = lines[0][:100] if lines and lines[0].strip() else file_path.stem

            return {
                "title": title,
                "content": content.strip(),
                "author": "Unknown",
                "language": "en",  # TODO: Detect language
                "word_count": len(content.split()),
                "line_count": len(lines),
            }

        except Exception as e:
            logger.error(f"Failed to extract text content from {file_path}: {e}")
            raise ValueError(f"Could not process text file: {str(e)}")

    async def _extract_web_content(self, url: str) -> Dict[str, Any]:
        """Extract content from web article"""
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=30) as response:
                    html_content = await response.text()

            # TODO: Implement HTML parsing and content extraction
            # For now, return basic structure
            return {
                "title": f"Web Article from {urlparse(url).netloc}",
                "content": "Web content extraction not yet implemented",
                "author": "Unknown",
                "language": "en",
                "word_count": 0,
                "url": url,
            }

        except Exception as e:
            logger.error(f"Failed to extract web content from {url}: {e}")
            raise ValueError(f"Could not process web article: {str(e)}")

    async def _analyze_content(
        self, content: str, metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Analyze content using AI to extract topics, difficulty, etc."""

        # Prepare analysis prompt
        analysis_prompt = f"""
        Analyze the following educational content and provide a structured analysis:

        Content Title: {metadata.get("title", "Unknown")}
        Content Length: {len(content.split())} words

        Content:
        {content[:2000]}...

        Please provide:
        1. Main topics covered (as a list)
        2. Difficulty level (beginner/intermediate/advanced)
        3. Key concepts (top 5-10 concepts)
        4. Estimated study time in minutes
        5. Language detected
        6. Content type classification (educational, technical, creative, etc.)

        Respond in JSON format:
        {{
            "topics": ["topic1", "topic2", ...],
            "difficulty_level": "beginner|intermediate|advanced",
            "key_concepts": ["concept1", "concept2", ...],
            "estimated_study_time": number_in_minutes,
            "language": "language_code",
            "content_classification": "educational|technical|creative|general"
        }}
        """

        try:
            # Use Ollama specifically for structured JSON responses (most reliable for JSON)
            messages = [{"role": "user", "content": analysis_prompt}]
            ollama_service = ai_router.providers.get("ollama")

            if ollama_service and ollama_service.is_available:
                response = await ollama_service.generate_response(
                    messages=messages,
                    language="en",
                    temperature=0.1,  # Low temperature for structured responses
                    max_tokens=500,
                )
            else:
                # Fallback to router if Ollama unavailable
                logger.warning(
                    "Ollama unavailable, falling back to AI router for content analysis"
                )
                response = await generate_ai_response(messages, language="en")

            # Parse JSON response
            analysis_result = json.loads(response.content)

            return {
                "topics": analysis_result.get("topics", []),
                "difficulty_level": analysis_result.get(
                    "difficulty_level", "intermediate"
                ),
                "key_concepts": analysis_result.get("key_concepts", []),
                "estimated_study_time": analysis_result.get("estimated_study_time", 30),
                "detected_language": analysis_result.get("language", "en"),
                "content_classification": analysis_result.get(
                    "content_classification", "general"
                ),
            }

        except Exception as e:
            logger.error(f"Content analysis failed: {e}")
            # Return default analysis
            return {
                "topics": ["General"],
                "difficulty_level": "intermediate",
                "key_concepts": [],
                "estimated_study_time": max(
                    len(content.split()) // 200, 5
                ),  # Rough estimate
                "detected_language": "en",
                "content_classification": "general",
            }

    async def _generate_learning_materials(
        self,
        content: str,
        metadata: ContentMetadata,
        material_types: List[LearningMaterialType],
    ) -> List[LearningMaterial]:
        """Generate various learning materials from content"""

        materials = []

        for material_type in material_types:
            try:
                material = await self._generate_single_material(
                    content, metadata, material_type
                )
                if material:
                    materials.append(material)
            except Exception as e:
                logger.error(f"Failed to generate {material_type.value}: {e}")

        return materials

    async def _generate_single_material(
        self,
        content: str,
        metadata: ContentMetadata,
        material_type: LearningMaterialType,
    ) -> Optional[LearningMaterial]:
        """Generate a single learning material of specified type"""

        # Prompts for different material types
        prompts = {
            LearningMaterialType.SUMMARY: f"""
                Create a comprehensive summary of the following content:

                Title: {metadata.title}

                Content:
                {content[:3000]}

                Provide a well-structured summary with:
                1. Main points (3-5 bullet points)
                2. Key takeaways
                3. Important details

                Format as JSON:
                {{
                    "main_points": ["point1", "point2", ...],
                    "key_takeaways": ["takeaway1", "takeaway2", ...],
                    "summary_text": "detailed summary paragraph"
                }}
            """,
            LearningMaterialType.FLASHCARDS: f"""
                Create 10-15 flashcards from the following content:

                Title: {metadata.title}

                Content:
                {content[:3000]}

                Create flashcards that test key concepts, definitions, and important facts.

                Format as JSON:
                {{
                    "flashcards": [
                        {{"front": "question or term", "back": "answer or definition"}},
                        ...
                    ]
                }}
            """,
            LearningMaterialType.QUIZ: f"""
                Create a 10-question quiz from the following content:

                Title: {metadata.title}

                Content:
                {content[:3000]}

                Include multiple choice, true/false, and short answer questions.

                Format as JSON:
                {{
                    "questions": [
                        {{
                            "type": "multiple_choice|true_false|short_answer",
                            "question": "question text",
                            "options": ["option1", "option2", ...] (for multiple choice only),
                            "correct_answer": "correct answer",
                            "explanation": "explanation of answer"
                        }},
                        ...
                    ]
                }}
            """,
            LearningMaterialType.KEY_CONCEPTS: f"""
                Extract and explain the key concepts from the following content:

                Title: {metadata.title}

                Content:
                {content[:3000]}

                Identify 8-12 key concepts and provide clear explanations.

                Format as JSON:
                {{
                    "concepts": [
                        {{
                            "term": "concept name",
                            "definition": "clear definition",
                            "importance": "why this concept matters",
                            "examples": ["example1", "example2", ...]
                        }},
                        ...
                    ]
                }}
            """,
            LearningMaterialType.NOTES: f"""
                Create structured study notes from the following content:

                Title: {metadata.title}

                Content:
                {content[:3000]}

                Organize into clear sections with bullet points and sub-points.

                Format as JSON:
                {{
                    "sections": [
                        {{
                            "title": "section title",
                            "content": ["bullet point 1", "bullet point 2", ...],
                            "subsections": [
                                {{
                                    "title": "subsection title",
                                    "content": ["detail 1", "detail 2", ...]
                                }}
                            ]
                        }},
                        ...
                    ]
                }}
            """,
        }

        prompt = prompts.get(material_type)
        if not prompt:
            logger.warning(f"No prompt defined for material type: {material_type}")
            return None

        try:
            # Use Ollama specifically for structured JSON responses (most reliable for JSON)
            messages = [{"role": "user", "content": prompt}]
            ollama_service = ai_router.providers.get("ollama")

            if ollama_service and ollama_service.is_available:
                response = await ollama_service.generate_response(
                    messages=messages,
                    language=metadata.language,
                    temperature=0.1,  # Low temperature for structured responses
                    max_tokens=800,
                )
            else:
                # Fallback to router if Ollama unavailable
                logger.warning(
                    "Ollama unavailable, falling back to AI router for material generation"
                )
                response = await generate_ai_response(
                    messages, language=metadata.language
                )

            # Parse JSON response
            material_content = json.loads(response.content)

            # Create learning material
            material_id = f"{metadata.content_id}_{material_type.value}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            return LearningMaterial(
                material_id=material_id,
                content_id=metadata.content_id,
                material_type=material_type,
                title=f"{material_type.value.title()} - {metadata.title}",
                content=material_content,
                difficulty_level=metadata.difficulty_level,
                estimated_time=self._estimate_material_time(
                    material_type, material_content
                ),
                tags=metadata.topics,
            )

        except Exception as e:
            logger.error(f"Failed to generate {material_type.value}: {e}")
            return None

    def _estimate_material_time(
        self, material_type: LearningMaterialType, content: Dict[str, Any]
    ) -> int:
        """Estimate time needed to study material (in minutes)"""

        estimates = {
            LearningMaterialType.SUMMARY: 5,
            LearningMaterialType.FLASHCARDS: lambda c: len(c.get("flashcards", []))
            * 0.5,  # 30 sec per card
            LearningMaterialType.QUIZ: lambda c: len(c.get("questions", []))
            * 1.5,  # 1.5 min per question
            LearningMaterialType.KEY_CONCEPTS: lambda c: len(c.get("concepts", []))
            * 2,  # 2 min per concept
            LearningMaterialType.NOTES: 10,
            LearningMaterialType.PRACTICE_QUESTIONS: lambda c: len(
                c.get("questions", [])
            )
            * 2,
        }

        estimate = estimates.get(material_type, 10)
        if callable(estimate):
            return max(int(estimate(content)), 1)
        return estimate

    async def process_content(
        self,
        source: str,
        file_path: Optional[Path] = None,
        material_types: Optional[List[LearningMaterialType]] = None,
        language: str = "en",
    ) -> str:
        """
        Process content and generate learning materials

        Args:
            source: URL or text content source
            file_path: Path to uploaded file (if applicable)
            material_types: Types of learning materials to generate
            language: Target language for processing

        Returns:
            Content ID for tracking progress
        """

        # Generate content ID
        content_id = self._generate_content_id(source)

        # Initialize progress tracking
        self._update_progress(
            content_id,
            ProcessingStatus.QUEUED,
            "Initializing",
            0,
            "Content processing started",
        )

        # Default material types
        if not material_types:
            material_types = [
                LearningMaterialType.SUMMARY,
                LearningMaterialType.FLASHCARDS,
                LearningMaterialType.KEY_CONCEPTS,
            ]

        try:
            # Start processing in background
            asyncio.create_task(
                self._process_content_async(
                    content_id, source, file_path, material_types, language
                )
            )

            return content_id

        except Exception as e:
            self._update_progress(
                content_id,
                ProcessingStatus.FAILED,
                "Initialization failed",
                0,
                f"Failed to start processing: {str(e)}",
                str(e),
            )
            raise

    async def _process_content_async(
        self,
        content_id: str,
        source: str,
        file_path: Optional[Path],
        material_types: List[LearningMaterialType],
        language: str,
    ):
        """Async content processing workflow"""

        try:
            start_time = datetime.now()

            # Step 1: Content Extraction (20%)
            self._update_progress(
                content_id,
                ProcessingStatus.EXTRACTING,
                "Extracting content",
                10,
                "Detecting content type and extracting data",
            )

            content_type = self._detect_content_type(source, file_path)

            # Extract content based on type
            if content_type == ContentType.YOUTUBE_VIDEO:
                raw_data = await self._extract_youtube_content(source)
            elif content_type == ContentType.PDF_DOCUMENT:
                raw_data = await self._extract_pdf_content(file_path)
            elif content_type == ContentType.WORD_DOCUMENT:
                raw_data = await self._extract_docx_content(file_path)
            elif content_type == ContentType.TEXT_FILE:
                raw_data = await self._extract_text_content(file_path)
            elif content_type == ContentType.WEB_ARTICLE:
                raw_data = await self._extract_web_content(source)
            else:
                raise ValueError(f"Unsupported content type: {content_type}")

            self._update_progress(
                content_id, ProcessingStatus.EXTRACTING, "Content extracted", 20
            )

            # Step 2: Content Analysis (40%)
            self._update_progress(
                content_id,
                ProcessingStatus.ANALYZING,
                "Analyzing content",
                30,
                "Analyzing content structure and topics",
            )

            content_text = raw_data["content"]
            analysis_result = await self._analyze_content(content_text, raw_data)

            # Create metadata
            metadata = ContentMetadata(
                content_id=content_id,
                title=raw_data["title"],
                content_type=content_type,
                source_url=source if not file_path else None,
                language=analysis_result.get("detected_language", language),
                duration=raw_data.get("duration"),
                word_count=raw_data["word_count"],
                difficulty_level=analysis_result["difficulty_level"],
                topics=analysis_result["topics"],
                author=raw_data.get("author"),
                created_at=datetime.now(),
                file_size=file_path.stat().st_size if file_path else None,
            )

            self._update_progress(
                content_id, ProcessingStatus.ANALYZING, "Analysis complete", 40
            )

            # Step 3: Learning Material Generation (80%)
            self._update_progress(
                content_id,
                ProcessingStatus.GENERATING,
                "Generating learning materials",
                50,
                f"Creating {len(material_types)} types of learning materials",
            )

            learning_materials = await self._generate_learning_materials(
                content_text, metadata, material_types
            )

            self._update_progress(
                content_id, ProcessingStatus.GENERATING, "Materials generated", 80
            )

            # Step 4: Organization and Storage (100%)
            self._update_progress(
                content_id,
                ProcessingStatus.ORGANIZING,
                "Organizing content",
                90,
                "Storing processed content and materials",
            )

            # Create processed content object
            processing_time = (datetime.now() - start_time).total_seconds()

            processed_content = ProcessedContent(
                metadata=metadata,
                raw_content=content_text,
                processed_content=content_text[
                    : self.max_content_length
                ],  # Truncate if needed
                learning_materials=learning_materials,
                processing_stats={
                    "processing_time": processing_time,
                    "materials_generated": len(learning_materials),
                    "content_length": len(content_text),
                    "language_detected": analysis_result.get("detected_language"),
                    "ai_analysis_topics": len(analysis_result["topics"]),
                },
            )

            # Store in content library
            self.content_library[content_id] = processed_content

            # Complete processing
            self._update_progress(
                content_id,
                ProcessingStatus.COMPLETED,
                "Processing completed",
                100,
                f"Successfully processed in {processing_time:.1f} seconds",
            )

            logger.info(
                f"Content {content_id} processed successfully in {processing_time:.1f}s"
            )

        except Exception as e:
            self._update_progress(
                content_id,
                ProcessingStatus.FAILED,
                "Processing failed",
                0,
                f"Error during processing: {str(e)}",
                str(e),
            )
            logger.error(f"Content processing failed for {content_id}: {e}")

    async def get_processed_content(
        self, content_id: str
    ) -> Optional[ProcessedContent]:
        """Get processed content by ID"""
        return self.content_library.get(content_id)

    async def get_content_library(self) -> List[Dict[str, Any]]:
        """Get all processed content metadata"""
        library = []
        for content_id, processed in self.content_library.items():
            library.append(
                {
                    "content_id": content_id,
                    "title": processed.metadata.title,
                    "content_type": processed.metadata.content_type.value,
                    "topics": processed.metadata.topics,
                    "difficulty_level": processed.metadata.difficulty_level,
                    "created_at": processed.metadata.created_at.isoformat(),
                    "material_count": len(processed.learning_materials),
                    "word_count": processed.metadata.word_count,
                    "estimated_study_time": sum(
                        material.estimated_time
                        for material in processed.learning_materials
                    ),
                }
            )

        # Sort by creation date (newest first)
        library.sort(key=lambda x: x["created_at"], reverse=True)
        return library

    async def search_content(
        self, query: str, filters: Optional[Dict[str, Any]] = None
    ) -> List[Dict[str, Any]]:
        """Search processed content"""
        results = []
        query_lower = query.lower()

        for content_id, processed in self.content_library.items():
            # Search in title, topics, and content
            matches = (
                query_lower in processed.metadata.title.lower()
                or any(
                    query_lower in topic.lower() for topic in processed.metadata.topics
                )
                or query_lower in processed.processed_content.lower()
            )

            if matches:
                # Apply filters if provided
                if filters:
                    if (
                        filters.get("content_type")
                        and processed.metadata.content_type.value
                        != filters["content_type"]
                    ):
                        continue
                    if (
                        filters.get("difficulty_level")
                        and processed.metadata.difficulty_level
                        != filters["difficulty_level"]
                    ):
                        continue
                    if (
                        filters.get("language")
                        and processed.metadata.language != filters["language"]
                    ):
                        continue

                results.append(
                    {
                        "content_id": content_id,
                        "title": processed.metadata.title,
                        "content_type": processed.metadata.content_type.value,
                        "topics": processed.metadata.topics,
                        "difficulty_level": processed.metadata.difficulty_level,
                        "relevance_score": self._calculate_relevance(query, processed),
                        "snippet": self._get_content_snippet(
                            query, processed.processed_content
                        ),
                    }
                )

        # Sort by relevance score
        results.sort(key=lambda x: x["relevance_score"], reverse=True)
        return results

    def _calculate_relevance(self, query: str, processed: ProcessedContent) -> float:
        """Calculate relevance score for search results"""
        score = 0.0
        query_lower = query.lower()

        # Title match (highest weight)
        if query_lower in processed.metadata.title.lower():
            score += 1.0

        # Topics match
        for topic in processed.metadata.topics:
            if query_lower in topic.lower():
                score += 0.5

        # Content match (lower weight)
        content_lower = processed.processed_content.lower()
        if query_lower in content_lower:
            score += 0.2

        return score

    def _get_content_snippet(
        self, query: str, content: str, max_length: int = 200
    ) -> str:
        """Get relevant snippet from content for search results"""
        query_lower = query.lower()
        content_lower = content.lower()

        # Find query position in content
        query_pos = content_lower.find(query_lower)

        if query_pos == -1:
            # Query not found, return beginning of content
            return (
                content[:max_length] + "..." if len(content) > max_length else content
            )

        # Extract snippet around query
        start = max(0, query_pos - max_length // 2)
        end = min(len(content), start + max_length)

        snippet = content[start:end]

        # Add ellipsis if needed
        if start > 0:
            snippet = "..." + snippet
        if end < len(content):
            snippet = snippet + "..."

        return snippet


# Global content processor instance
content_processor = ContentProcessor()
